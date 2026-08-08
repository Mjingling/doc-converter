#!/usr/bin/env python3
"""生成 DocMorph 手动验证测试素材（输出到 workspace 内 .manual-test/）"""
import os, shutil
from PIL import Image, ImageDraw, ImageFont

BASE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", ".manual-test")
BASE = os.path.abspath(BASE)
os.makedirs(BASE, exist_ok=True)

FONT_CANDIDATES = [
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
]
FONT = next((f for f in FONT_CANDIDATES if os.path.exists(f)), None)
print("font:", FONT)

def gradient(size, c1, c2, text, num):
    img = Image.new("RGB", size)
    d = ImageDraw.Draw(img)
    w, h = size
    for y in range(h):
        t = y / h
        d.line([(0, y), (w, y)], fill=tuple(int(c1[i] + (c2[i] - c1[i]) * t) for i in range(3)))
    f = ImageFont.truetype(FONT, 80) if FONT else ImageFont.load_default()
    d.text((80, 120), f"DocMorph Test Page {num}", font=f, fill=(255, 255, 255))
    d.text((80, 240), text, font=f, fill=(255, 255, 255))
    for i in range(20):
        d.ellipse([(i * 137 % w, i * 211 % h), (i * 137 % w + 90, i * 211 % h + 90)],
                  outline=(255, 255, 255, 128), width=6)
    return img

# 1) 大图两张（合成多页大 PDF）
big1 = gradient((2000, 1500), (30, 60, 120), (200, 60, 60), "中文渐变页 A", 1)
big2 = gradient((2000, 1500), (20, 100, 80), (60, 60, 200), "中文渐变页 B", 2)
big1.save(f"{BASE}/big1.png")
big2.save(f"{BASE}/big2.png")

# 2) 多页 PDF（图片型，体积大，用于验证水印/旋转/页码/加密）
big1.save(f"{BASE}/sample.pdf", save_all=True, append_images=[big2], resolution=150)
print("sample.pdf size:", os.path.getsize(f"{BASE}/sample.pdf"))

# 3) 小图两张（用于图片转 PDF）
for i, (c1, c2) in enumerate([((240, 120, 60), (60, 180, 240)), ((60, 200, 120), (180, 60, 220))]):
    gradient((800, 600), c1, c2, f"Image {i + 1}", i + 1).save(f"{BASE}/img{i + 1}.png")

# 4) docx（用于 ConvertPanel 转换 + 文件夹批量）
import docx
doc = docx.Document()
doc.add_heading("DocMorph 转换测试文档", 0)
for i in range(8):
    doc.add_paragraph(f"这是第 {i + 1} 段测试内容：DocMorph 是一个办公文档转换与 PDF 工具箱桌面应用。")
    doc.add_paragraph(f"This is paragraph {i + 1} for conversion testing.")
doc.save(f"{BASE}/sample.docx")

# 5) 文件夹批量素材：混合格式
BATCH = f"{BASE}/batch"
os.makedirs(BATCH, exist_ok=True)
shutil.copy(f"{BASE}/sample.docx", f"{BATCH}/a.docx")
shutil.copy(f"{BASE}/sample.pdf", f"{BATCH}/b.pdf")
shutil.copy(f"{BASE}/img1.png", f"{BATCH}/c.png")
with open(f"{BATCH}/d.md", "w") as f:
    f.write("# Markdown 测试\n\n这是一段 markdown 内容。\n")
with open(f"{BATCH}/e.txt", "w") as f:
    f.write("纯文本测试内容。\n")
print("done, files:", sorted(os.listdir(BASE)))
